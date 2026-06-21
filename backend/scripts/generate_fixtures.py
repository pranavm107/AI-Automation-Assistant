import os
import sys

sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from pypdf import PdfWriter
from docx import Document

def generate():
    fixtures_dir = "tests/fixtures"
    os.makedirs(fixtures_dir, exist_ok=True)

    # 1. Empty TXT
    open(f"{fixtures_dir}/empty.txt", "w").close()

    # 2. Sample TXT
    with open(f"{fixtures_dir}/sample_notes.txt", "w") as f:
        f.write("This is a sample text document.\nIt has multiple lines.\n\n\n\nAnd some extra     spacing.")

    # 3. Sample DOCX
    doc = Document()
    doc.add_paragraph("This is a sample DOCX document.")
    doc.add_paragraph("It contains multiple paragraphs.")
    doc.save(f"{fixtures_dir}/sample_document.docx")

    # 4. Empty PDF
    writer = PdfWriter()
    writer.add_blank_page(width=72, height=72)
    with open(f"{fixtures_dir}/sample_resume.pdf", "wb") as f:
        writer.write(f)
        
    # 5. Corrupted PDF
    with open(f"{fixtures_dir}/corrupted.pdf", "w") as f:
        f.write("This is not a real PDF file")
        
    print("Fixtures generated successfully.")

if __name__ == "__main__":
    generate()
